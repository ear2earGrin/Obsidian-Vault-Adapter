#!/usr/bin/env python3
"""
vault_builder.py

Watches a source folder, extracts text from every PDF/DOCX it finds,
enriches each document with a local Qwen3 model via LM Studio, and writes
clean Markdown notes with YAML frontmatter into an Obsidian vault.

Usage:
    python vault_builder.py --source /path/to/docs --vault /path/to/vault
    python vault_builder.py --config /path/to/config.yaml
"""

import argparse
import hashlib
import json
import logging
import re
import sys
import time
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import Optional

import requests
import yaml
from rich.console import Console
from rich.logging import RichHandler
from rich.progress import (
    BarColumn,
    MofNCompleteColumn,
    Progress,
    SpinnerColumn,
    TaskProgressColumn,
    TextColumn,
    TimeElapsedColumn,
)
from rich.table import Table

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------

logging.basicConfig(
    level=logging.INFO,
    format="%(message)s",
    datefmt="[%X]",
    handlers=[RichHandler(rich_tracebacks=True)],
)
log = logging.getLogger("vault_builder")
console = Console()

# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------


@dataclass
class ExtractedDoc:
    title: str
    raw_text: str
    headings: list[str]
    metadata: dict
    page_count: int
    word_count: int
    source_path: Path
    file_type: str  # "pdf" or "docx"


@dataclass
class EnrichedDoc:
    extracted: ExtractedDoc
    summary: str = ""
    tags: list[str] = field(default_factory=list)
    inferred_title: str = ""
    key_concepts: list[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------


def load_config(config_path: Optional[Path], source: Optional[str], vault: Optional[str]) -> dict:
    defaults = {
        "source_path": source,
        "vault_path": vault,
        "lm_studio": {
            "endpoint": "http://localhost:1234/v1/chat/completions",
            "model": "qwen3",
            "timeout_seconds": 120,
            "max_retries": 3,
            "retry_delay_seconds": 5,
        },
        "batch_size": 10,
        "enrichment_word_limit": 2000,
        "split_word_threshold": 5000,
        "folders": {
            "inbox": "Inbox",
            "pdfs": "Documents/PDFs",
            "word": "Documents/Word",
            "chatgpt": "Documents/ChatGPT",
            "moc": "MOC",
        },
        "state_file": "processed.json",
    }

    if config_path and config_path.exists():
        with open(config_path) as f:
            file_cfg = yaml.safe_load(f) or {}
        # CLI flags override config file
        if source:
            file_cfg["source_path"] = source
        if vault:
            file_cfg["vault_path"] = vault
        # Deep merge lm_studio sub-dict
        merged_lm = {**defaults["lm_studio"], **file_cfg.get("lm_studio", {})}
        merged_folders = {**defaults["folders"], **file_cfg.get("folders", {})}
        cfg = {**defaults, **file_cfg, "lm_studio": merged_lm, "folders": merged_folders}
    else:
        cfg = defaults

    if not cfg.get("vault_path"):
        log.error("vault_path is required. Use --vault or set it in config.yaml.")
        sys.exit(1)

    return cfg


# ---------------------------------------------------------------------------
# Stage 1 — File scanner + state tracker
# ---------------------------------------------------------------------------


def file_hash(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def load_state(state_file: Path) -> dict:
    if state_file.exists():
        with open(state_file) as f:
            return json.load(f)
    return {}


def save_state(state_file: Path, state: dict) -> None:
    state_file.parent.mkdir(parents=True, exist_ok=True)
    with open(state_file, "w") as f:
        json.dump(state, f, indent=2)


def scan_files(source_path: Path, state: dict) -> list[Path]:
    """Return files that haven't been processed yet (by hash)."""
    extensions = {".pdf", ".docx"}
    queue: list[Path] = []
    skipped = 0

    for path in sorted(source_path.rglob("*")):
        if path.suffix.lower() not in extensions:
            continue
        if not path.is_file():
            continue
        h = file_hash(path)
        if h in state:
            skipped += 1
            continue
        queue.append(path)

    if skipped:
        log.info(f"Skipped {skipped} already-processed file(s).")
    log.info(f"Found {len(queue)} new file(s) to process.")
    return queue


# ---------------------------------------------------------------------------
# Stage 2 — Extractors
# ---------------------------------------------------------------------------


def _clean_text(text: str) -> str:
    """Normalise whitespace without destroying paragraph breaks."""
    lines = text.splitlines()
    cleaned = []
    for line in lines:
        stripped = line.strip()
        cleaned.append(stripped)
    # Collapse runs of 3+ blank lines to 2
    result = re.sub(r"\n{3,}", "\n\n", "\n".join(cleaned))
    return result.strip()


def _stem_title(path: Path) -> str:
    """Best-effort human title from a filename."""
    stem = path.stem
    stem = re.sub(r"[_\-]+", " ", stem)
    stem = re.sub(r"\s+", " ", stem).strip()
    return stem.title()


def extract_pdf(path: Path) -> ExtractedDoc:
    try:
        import pdfplumber
    except ImportError:
        log.error("pdfplumber is not installed. Run: pip install pdfplumber")
        sys.exit(1)

    text_parts: list[str] = []
    headings: list[str] = []
    page_count = 0
    metadata: dict = {}

    try:
        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            if pdf.metadata:
                metadata = {k: str(v) for k, v in pdf.metadata.items() if v}

            for page in pdf.pages:
                # Extract words with font-size info to detect headings
                words = page.extract_words(extra_attrs=["size", "fontname"])
                if not words:
                    continue

                # Determine median font size on this page for heading heuristic
                sizes = [w.get("size", 0) for w in words if w.get("size")]
                if sizes:
                    median_size = sorted(sizes)[len(sizes) // 2]
                    heading_threshold = median_size * 1.3
                else:
                    heading_threshold = float("inf")

                page_text = page.extract_text() or ""
                text_parts.append(page_text)

                # Collect heading candidates
                current_heading_words: list[str] = []
                for w in words:
                    if w.get("size", 0) >= heading_threshold:
                        current_heading_words.append(w["text"])
                    else:
                        if current_heading_words:
                            candidate = " ".join(current_heading_words).strip()
                            if candidate and candidate not in headings:
                                headings.append(candidate)
                            current_heading_words = []

    except Exception as exc:
        log.warning(f"pdfplumber failed on {path.name}: {exc}. Attempting PyMuPDF fallback.")
        text_parts, headings, page_count = _extract_pdf_pymupdf(path)

    raw_text = _clean_text("\n\n".join(text_parts))
    if not raw_text:
        log.warning(f"{path.name}: no text layer found — may be a scanned PDF. Flagging for manual review.")
        raw_text = "[SCANNED PDF — no text layer detected. Manual OCR required.]"

    title = metadata.get("Title", "") or _stem_title(path)

    return ExtractedDoc(
        title=title,
        raw_text=raw_text,
        headings=headings[:20],  # cap to keep things sane
        metadata=metadata,
        page_count=page_count,
        word_count=len(raw_text.split()),
        source_path=path,
        file_type="pdf",
    )


def _extract_pdf_pymupdf(path: Path) -> tuple[list[str], list[str], int]:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        log.error("PyMuPDF is not installed. Run: pip install pymupdf")
        return [], [], 0

    text_parts: list[str] = []
    headings: list[str] = []
    try:
        doc = fitz.open(str(path))
        page_count = len(doc)
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
    except Exception as exc:
        log.error(f"PyMuPDF also failed on {path.name}: {exc}")
        return [], [], 0

    return text_parts, headings, page_count


def extract_docx(path: Path) -> ExtractedDoc:
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        log.error("python-docx is not installed. Run: pip install python-docx")
        sys.exit(1)

    text_parts: list[str] = []
    headings: list[str] = []
    metadata: dict = {}

    try:
        doc = Document(str(path))

        # Core properties
        cp = doc.core_properties
        metadata = {
            "author": cp.author or "",
            "created": str(cp.created or ""),
            "title": cp.title or "",
        }

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                continue

            if style_name.startswith("Heading"):
                # Map Heading 1 → #, Heading 2 → ##, etc.
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = 1
                prefix = "#" * min(level, 6)
                text_parts.append(f"{prefix} {text}")
                headings.append(text)
            else:
                text_parts.append(text)

        # Tables → Markdown pipe tables
        for table in doc.tables:
            rows = []
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    rows.append("|" + " --- |" * len(cells))
            text_parts.extend(rows)
            text_parts.append("")

    except Exception as exc:
        log.warning(f"python-docx failed on {path.name}: {exc}. Trying mammoth fallback.")
        text_parts, headings = _extract_docx_mammoth(path)

    raw_text = _clean_text("\n\n".join(text_parts))
    title = metadata.get("title", "") or _stem_title(path)

    return ExtractedDoc(
        title=title,
        raw_text=raw_text,
        headings=headings[:20],
        metadata=metadata,
        page_count=0,  # DOCX has no native page count
        word_count=len(raw_text.split()),
        source_path=path,
        file_type="docx",
    )


def _extract_docx_mammoth(path: Path) -> tuple[list[str], list[str]]:
    try:
        import mammoth
    except ImportError:
        log.error("mammoth is not installed. Run: pip install mammoth")
        return [], []

    try:
        with open(path, "rb") as f:
            result = mammoth.convert_to_markdown(f)
        text = result.value
        # Pull headings from markdown output
        headings = [line.lstrip("# ").strip() for line in text.splitlines() if line.startswith("#")]
        return [text], headings
    except Exception as exc:
        log.error(f"mammoth also failed on {path.name}: {exc}")
        return [], []


def extract(path: Path) -> ExtractedDoc:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path)
    elif suffix == ".docx":
        return extract_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


# ---------------------------------------------------------------------------
# Stage 2b — ChatGPT export parser
# ---------------------------------------------------------------------------


def _reconstruct_messages(mapping: dict) -> list[dict]:
    """Walk the conversation tree and return messages in order."""
    # Find root node (no parent, or parent not in mapping)
    root_id = None
    for node_id, node in mapping.items():
        parent = node.get("parent")
        if parent is None or parent not in mapping:
            root_id = node_id
            break

    if root_id is None:
        return []

    messages: list[dict] = []

    def traverse(node_id: str) -> None:
        node = mapping.get(node_id)
        if not node:
            return
        msg = node.get("message")
        if msg:
            role = msg.get("author", {}).get("role", "")
            content = msg.get("content", {})
            parts = content.get("parts", [])
            # Only keep plain text parts (skip image/tool blocks)
            text = "\n".join(str(p) for p in parts if isinstance(p, str)).strip()
            if text and role in ("user", "assistant"):
                messages.append({
                    "role": role,
                    "text": text,
                    "time": msg.get("create_time"),
                    "model": msg.get("metadata", {}).get("model_slug", ""),
                })
        for child_id in node.get("children", []):
            traverse(child_id)

    traverse(root_id)
    return messages


def parse_chatgpt_export(export_path: Path, min_words: int = 80) -> list[ExtractedDoc]:
    """
    Parse a ChatGPT conversations.json export.
    Returns one ExtractedDoc per conversation (skipping trivially short ones).
    """
    log.info(f"Parsing ChatGPT export: {export_path.name}")

    with open(export_path, encoding="utf-8") as f:
        data = json.load(f)

    if not isinstance(data, list):
        raise ValueError("conversations.json must be a JSON array at the top level")

    docs: list[ExtractedDoc] = []

    for conv in data:
        title = (conv.get("title") or "Untitled Conversation").strip()
        create_time = conv.get("create_time")
        mapping = conv.get("mapping") or {}

        messages = _reconstruct_messages(mapping)
        if not messages:
            continue

        # Format as readable dialogue
        lines: list[str] = []
        for m in messages:
            speaker = "**You**" if m["role"] == "user" else "**ChatGPT**"
            lines.append(f"{speaker}: {m['text']}\n")

        raw_text = "\n".join(lines)
        word_count = len(raw_text.split())

        if word_count < min_words:
            continue

        # Derive ISO date from timestamp
        conv_date = ""
        if create_time:
            from datetime import datetime, timezone
            try:
                conv_date = datetime.fromtimestamp(create_time, tz=timezone.utc).strftime("%Y-%m-%d")
            except Exception:
                pass

        # Use a stable unique ID as the "path" for state tracking
        conv_id = conv.get("conversation_id") or conv.get("id") or title
        # Fake path so the rest of the pipeline treats it like a file
        fake_path = export_path.parent / f"chatgpt__{conv_id}.chatgpt"

        metadata = {
            "conversation_id": str(conv_id),
            "date": conv_date,
            "model": messages[-1].get("model", "") if messages else "",
            "message_count": str(len(messages)),
        }

        docs.append(ExtractedDoc(
            title=title,
            raw_text=raw_text,
            headings=[],
            metadata=metadata,
            page_count=0,
            word_count=word_count,
            source_path=fake_path,
            file_type="chatgpt",
        ))

    log.info(f"Parsed {len(docs)} conversations (skipped {len(data) - len(docs)} too short)")
    return docs


def chatgpt_state_key(doc: ExtractedDoc) -> str:
    """Stable hash key for a ChatGPT conversation (based on conversation_id)."""
    conv_id = doc.metadata.get("conversation_id", doc.title)
    return hashlib.sha256(conv_id.encode()).hexdigest()


# ---------------------------------------------------------------------------
# Stage 3 — LM Studio AI enrichment
# ---------------------------------------------------------------------------

ENRICHMENT_PROMPT = """\
You are a knowledge-base assistant. Analyse the document excerpt below and return ONLY a JSON object — no prose, no code fences, just valid JSON.

The JSON must have exactly these keys:
- "summary": A 3-5 sentence plain-English summary of what this document is about.
- "tags": An array of 5-10 lowercase single-word or hyphenated tags (e.g. "bitcoin", "ai-security", "strategy").
- "inferred_title": A clean, human-readable title for this document (improve on the filename if it is cryptic).
- "key_concepts": An array of 3-8 noun phrases that represent the core concepts (e.g. "prompt injection", "hash-chained logs").

Document title: {title}
Document excerpt ({word_count} words):
---
{excerpt}
---
"""


def call_lm_studio(excerpt: str, title: str, word_count: int, cfg: dict) -> dict:
    lm = cfg["lm_studio"]
    prompt = ENRICHMENT_PROMPT.format(title=title, word_count=word_count, excerpt=excerpt)

    payload = {
        "model": lm["model"],
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.3,
        "max_tokens": 512,
    }

    for attempt in range(1, lm["max_retries"] + 1):
        try:
            resp = requests.post(
                lm["endpoint"],
                json=payload,
                timeout=lm["timeout_seconds"],
            )
            resp.raise_for_status()
            content = resp.json()["choices"][0]["message"]["content"].strip()

            # Strip markdown code fences if the model wraps JSON anyway
            content = re.sub(r"^```(?:json)?\s*", "", content)
            content = re.sub(r"\s*```$", "", content)

            return json.loads(content)

        except (requests.RequestException, json.JSONDecodeError, KeyError) as exc:
            log.warning(f"LM Studio attempt {attempt}/{lm['max_retries']} failed: {exc}")
            if attempt < lm["max_retries"]:
                time.sleep(lm["retry_delay_seconds"])

    log.error("All LM Studio retries exhausted. Using empty enrichment.")
    return {
        "summary": "",
        "tags": [],
        "inferred_title": title,
        "key_concepts": [],
    }


def enrich(doc: ExtractedDoc, cfg: dict) -> EnrichedDoc:
    word_limit = cfg["enrichment_word_limit"]
    words = doc.raw_text.split()
    excerpt = " ".join(words[:word_limit])

    result = call_lm_studio(
        excerpt=excerpt,
        title=doc.title,
        word_count=min(len(words), word_limit),
        cfg=cfg,
    )

    return EnrichedDoc(
        extracted=doc,
        summary=result.get("summary", ""),
        tags=[str(t).lower().strip().lstrip("#") for t in result.get("tags", [])],
        inferred_title=result.get("inferred_title", doc.title) or doc.title,
        key_concepts=[str(c).strip() for c in result.get("key_concepts", [])],
    )


# ---------------------------------------------------------------------------
# Stage 4 — Markdown assembler
# ---------------------------------------------------------------------------


def _inject_wikilinks(text: str, concepts: list[str]) -> str:
    """Replace concept phrases in body text with [[wikilink]] format."""
    for concept in concepts:
        # Case-insensitive whole-phrase match, not already inside brackets
        pattern = r"(?<!\[)\b(" + re.escape(concept) + r")\b(?!\])"
        replacement = r"[[\1]]"
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
    return text


def _safe_filename(name: str) -> str:
    """Strip characters that are unsafe in filenames."""
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "", name)
    name = name.strip(". ")
    return name[:200] or "untitled"


def assemble_markdown(enriched: EnrichedDoc) -> str:
    doc = enriched.extracted
    today = date.today().isoformat()

    frontmatter = {
        "title": enriched.inferred_title or doc.title,
        "source": doc.source_path.name,
        "file_type": doc.file_type,
        "processed_date": today,
        "tags": enriched.tags,
        "summary": enriched.summary,
        "word_count": doc.word_count,
        "page_count": doc.page_count,
        "key_concepts": enriched.key_concepts,
    }

    fm_lines = ["---"]
    for k, v in frontmatter.items():
        if isinstance(v, list):
            if v:
                fm_lines.append(f"{k}:")
                for item in v:
                    fm_lines.append(f'  - "{item}"')
            else:
                fm_lines.append(f"{k}: []")
        elif isinstance(v, str) and ("\n" in v or '"' in v):
            # Block scalar for multi-line strings
            fm_lines.append(f"{k}: |")
            for line in v.splitlines():
                fm_lines.append(f"  {line}")
        else:
            safe_v = str(v).replace('"', '\\"')
            fm_lines.append(f'{k}: "{safe_v}"')
    fm_lines.append("---")
    fm_lines.append("")

    sections: list[str] = []

    # Summary section
    if enriched.summary:
        sections.append("## Summary\n")
        sections.append(enriched.summary)
        sections.append("\n---\n")

    # Document body with wikilinks injected
    body = _inject_wikilinks(doc.raw_text, enriched.key_concepts)
    sections.append(body)

    return "\n".join(fm_lines) + "\n".join(sections)


def _split_by_headings(text: str) -> list[tuple[str, str]]:
    """
    Split a Markdown body by top-level (#) headings.
    Returns list of (heading_title, section_content) tuples.
    First tuple uses 'Introduction' as heading if no leading # heading.
    """
    parts: list[tuple[str, str]] = []
    current_title = "Introduction"
    current_lines: list[str] = []

    for line in text.splitlines():
        if line.startswith("# ") and not line.startswith("## "):
            if current_lines:
                parts.append((current_title, "\n".join(current_lines).strip()))
            current_title = line.lstrip("# ").strip()
            current_lines = []
        else:
            current_lines.append(line)

    if current_lines:
        parts.append((current_title, "\n".join(current_lines).strip()))

    return parts


# ---------------------------------------------------------------------------
# Stage 5 — Vault output + MOC generation
# ---------------------------------------------------------------------------


def ensure_vault_structure(vault: Path, folders: dict) -> dict[str, Path]:
    paths: dict[str, Path] = {}
    for key, rel in folders.items():
        p = vault / rel
        p.mkdir(parents=True, exist_ok=True)
        paths[key] = p
    return paths


def write_note(
    enriched: EnrichedDoc,
    vault_paths: dict[str, Path],
    cfg: dict,
) -> list[Path]:
    """Write note(s) to vault, splitting large docs by heading. Returns written paths."""
    doc = enriched.extracted
    folder_key = {"pdf": "pdfs", "docx": "word", "chatgpt": "chatgpt"}.get(doc.file_type, "word")
    dest_folder = vault_paths[folder_key]

    base_name = _safe_filename(enriched.inferred_title or doc.title)
    written: list[Path] = []

    if doc.word_count > cfg["split_word_threshold"]:
        # Split by top-level headings
        body = _inject_wikilinks(doc.raw_text, enriched.key_concepts)
        sections = _split_by_headings(body)

        # Write a parent index note
        parent_content = _build_parent_index(enriched, base_name, sections)
        parent_path = dest_folder / f"{base_name}.md"
        parent_path.write_text(parent_content, encoding="utf-8")
        written.append(parent_path)

        # Write each section
        for heading, section_body in sections:
            section_name = _safe_filename(f"{base_name} — {heading}")
            section_path = dest_folder / f"{section_name}.md"
            section_content = _build_section_note(enriched, heading, section_body, base_name)
            section_path.write_text(section_content, encoding="utf-8")
            written.append(section_path)
    else:
        content = assemble_markdown(enriched)
        note_path = dest_folder / f"{base_name}.md"
        note_path.write_text(content, encoding="utf-8")
        written.append(note_path)

    return written


def _build_parent_index(enriched: EnrichedDoc, base_name: str, sections: list[tuple[str, str]]) -> str:
    doc = enriched.extracted
    today = date.today().isoformat()

    fm = [
        "---",
        f'title: "{enriched.inferred_title or doc.title}"',
        f'source: "{doc.source_path.name}"',
        f'file_type: "{doc.file_type}"',
        f'processed_date: "{today}"',
        f"tags: [{', '.join(repr(t) for t in enriched.tags)}]",
        f'word_count: "{doc.word_count}"',
        f'page_count: "{doc.page_count}"',
        "type: index",
        "---",
        "",
    ]

    lines = fm + [
        f"## {enriched.inferred_title or doc.title}",
        "",
        f"> {enriched.summary}" if enriched.summary else "",
        "",
        "### Sections",
        "",
    ]

    for heading, _ in sections:
        section_name = _safe_filename(f"{base_name} — {heading}")
        lines.append(f"- [[{section_name}]]")

    return "\n".join(lines)


def _build_section_note(enriched: EnrichedDoc, heading: str, body: str, parent_name: str) -> str:
    doc = enriched.extracted
    today = date.today().isoformat()

    fm = [
        "---",
        f'title: "{heading}"',
        f'parent: "[[{parent_name}]]"',
        f'source: "{doc.source_path.name}"',
        f'processed_date: "{today}"',
        f"tags: [{', '.join(repr(t) for t in enriched.tags)}]",
        "---",
        "",
        f"# {heading}",
        "",
        f"*Part of [[{parent_name}]]*",
        "",
    ]

    return "\n".join(fm) + body


def generate_moc_all(enriched_docs: list[EnrichedDoc], moc_folder: Path) -> None:
    today = date.today().isoformat()
    lines = [
        "---",
        f'title: "MOC — All Documents"',
        f'generated: "{today}"',
        "---",
        "",
        "# Master Index — All Documents",
        "",
        f"*Generated {today} — {len(enriched_docs)} document(s)*",
        "",
        "| Title | Source | Tags | Summary |",
        "| ----- | ------ | ---- | ------- |",
    ]

    for e in sorted(enriched_docs, key=lambda x: x.inferred_title.lower()):
        title = (e.inferred_title or e.extracted.title).replace("|", "\\|")
        source = e.extracted.source_path.name.replace("|", "\\|")
        tags = ", ".join(f"`{t}`" for t in e.tags)
        summary = (e.summary or "").replace("\n", " ").replace("|", "\\|")[:120]
        safe_title = _safe_filename(title)
        lines.append(f"| [[{safe_title}]] | {source} | {tags} | {summary} |")

    moc_path = moc_folder / "MOC-All.md"
    moc_path.write_text("\n".join(lines), encoding="utf-8")
    log.info(f"Wrote {moc_path}")


def generate_moc_tags(enriched_docs: list[EnrichedDoc], moc_folder: Path) -> None:
    today = date.today().isoformat()

    # Group by tag
    tag_map: dict[str, list[EnrichedDoc]] = {}
    for e in enriched_docs:
        for tag in e.tags:
            tag_map.setdefault(tag, []).append(e)

    lines = [
        "---",
        f'title: "MOC — By Tag"',
        f'generated: "{today}"',
        "---",
        "",
        "# Index by Tag",
        "",
        f"*Generated {today}*",
        "",
    ]

    for tag in sorted(tag_map.keys()):
        lines.append(f"## #{tag}")
        lines.append("")
        for e in sorted(tag_map[tag], key=lambda x: x.inferred_title.lower()):
            title = e.inferred_title or e.extracted.title
            safe_title = _safe_filename(title)
            lines.append(f"- [[{safe_title}]]")
        lines.append("")

    moc_path = moc_folder / "MOC-Tags.md"
    moc_path.write_text("\n".join(lines), encoding="utf-8")
    log.info(f"Wrote {moc_path}")


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------


def build_progress() -> Progress:
    return Progress(
        SpinnerColumn(),
        TextColumn("[bold blue]{task.description}"),
        BarColumn(),
        MofNCompleteColumn(),
        TaskProgressColumn(),
        TimeElapsedColumn(),
        console=console,
    )


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Convert PDFs and DOCXs into an AI-readable Obsidian vault.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("--source", help="Folder containing source documents (PDFs/DOCXs)")
    parser.add_argument("--chatgpt", metavar="PATH", help="Path to ChatGPT conversations.json export")
    parser.add_argument("--vault", help="Obsidian vault root folder")
    parser.add_argument("--config", default="config.yaml", help="Path to config.yaml (default: config.yaml)")
    parser.add_argument("--dry-run", action="store_true", help="Scan and extract only — skip AI enrichment and writing")
    parser.add_argument("--verbose", action="store_true", help="Show DEBUG-level logging")
    args = parser.parse_args()

    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    cfg = load_config(
        config_path=Path(args.config) if args.config else None,
        source=args.source,
        vault=args.vault,
    )

    vault_path = Path(cfg["vault_path"]).expanduser().resolve()
    vault_path.mkdir(parents=True, exist_ok=True)
    state_file = vault_path / cfg["state_file"]
    state = load_state(state_file)

    console.rule("[bold]Vault Builder[/bold]")
    console.print(f"  Vault  : [cyan]{vault_path}[/cyan]")
    console.print(f"  Model  : [cyan]{cfg['lm_studio']['model']}[/cyan] @ {cfg['lm_studio']['endpoint']}")
    console.print()

    vault_paths = ensure_vault_structure(vault_path, cfg["folders"])
    all_enriched: list[EnrichedDoc] = []

    # ── Build work queue ────────────────────────────────────────────────────
    # Either a list of file paths (PDF/DOCX mode) or pre-extracted docs (ChatGPT mode)
    chatgpt_docs: list[ExtractedDoc] = []

    if args.chatgpt:
        chatgpt_path = Path(args.chatgpt).expanduser().resolve()
        if not chatgpt_path.exists():
            log.error(f"ChatGPT export not found: {chatgpt_path}")
            sys.exit(1)
        console.print(f"  Source : [cyan]{chatgpt_path}[/cyan] (ChatGPT export)")
        all_convs = parse_chatgpt_export(chatgpt_path)
        # Filter already-processed conversations
        chatgpt_docs = [d for d in all_convs if chatgpt_state_key(d) not in state]
        log.info(f"{len(chatgpt_docs)} new conversations to process ({len(all_convs) - len(chatgpt_docs)} already done).")
        queue: list[Path] = []
    else:
        if not cfg.get("source_path"):
            log.error("source_path is required unless using --chatgpt.")
            sys.exit(1)
        source_path = Path(cfg["source_path"]).expanduser().resolve()
        if not source_path.exists():
            log.error(f"Source path does not exist: {source_path}")
            sys.exit(1)
        console.print(f"  Source : [cyan]{source_path}[/cyan]")
        queue = scan_files(source_path, state)
        if not queue:
            console.print("[green]Nothing new to process. Vault is up to date.[/green]")
            return

    total_items = len(chatgpt_docs) if args.chatgpt else len(queue)
    if total_items == 0:
        console.print("[green]Nothing new to process. Vault is up to date.[/green]")
        return

    batch_size = cfg["batch_size"]

    # ── Process ─────────────────────────────────────────────────────────────
    with build_progress() as progress:
        overall = progress.add_task("Overall", total=total_items)

        items = chatgpt_docs if args.chatgpt else queue

        for batch_start in range(0, total_items, batch_size):
            batch = items[batch_start : batch_start + batch_size]

            for item in batch:
                # item is ExtractedDoc for ChatGPT, Path for PDF/DOCX
                if args.chatgpt:
                    doc = item  # already extracted
                    label = doc.title[:60]
                else:
                    file_path = item
                    progress.update(overall, description=f"[bold]{file_path.name}[/bold]")
                    try:
                        doc = extract(file_path)
                    except Exception as exc:
                        log.error(f"Extraction failed for {file_path.name}: {exc}")
                        progress.advance(overall)
                        continue
                    label = file_path.name

                progress.update(overall, description=f"[bold]{label[:60]}[/bold]")

                if args.dry_run:
                    console.print(f"  [dim]dry-run[/dim] {label}: {doc.word_count} words")
                    progress.advance(overall)
                    continue

                # Stage 3 — Enrich
                try:
                    enriched = enrich(doc, cfg)
                except Exception as exc:
                    log.error(f"Enrichment failed for {label}: {exc}")
                    enriched = EnrichedDoc(extracted=doc, inferred_title=doc.title)

                # Stage 4+5 — Assemble and write
                try:
                    written_paths = write_note(enriched, vault_paths, cfg)
                    all_enriched.append(enriched)

                    state_key = chatgpt_state_key(doc) if args.chatgpt else file_hash(item)
                    state[state_key] = {
                        "source": str(doc.source_path),
                        "notes": [str(p) for p in written_paths],
                        "processed_date": date.today().isoformat(),
                    }
                    save_state(state_file, state)

                except Exception as exc:
                    log.error(f"Failed to write note for {label}: {exc}")

                progress.advance(overall)

    if not args.dry_run and all_enriched:
        # Regenerate MOC files
        moc_folder = vault_paths["moc"]
        generate_moc_all(all_enriched, moc_folder)
        generate_moc_tags(all_enriched, moc_folder)

    # Summary table
    console.print()
    table = Table(title="Run Summary", show_header=True, header_style="bold magenta")
    table.add_column("Metric", style="dim")
    table.add_column("Value", justify="right")
    table.add_row("Files found", str(len(queue)))
    table.add_row("Notes written", str(len(all_enriched)))
    table.add_row("Total processed (all time)", str(len(state)))
    console.print(table)
    console.print("[green]Done.[/green]")


if __name__ == "__main__":
    main()
